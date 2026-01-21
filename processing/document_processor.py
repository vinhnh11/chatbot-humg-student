import os
import re
import logging
from docx import Document as DocxDocument
from langchain_core.documents import Document

from config import DOCX_DIR

logger = logging.getLogger(__name__)


class DocumentProcessor:
    def __init__(self, verbose: bool = False):
        """
        verbose = True  -> in log khi debug / crawl
        verbose = False -> im lặng khi chạy chatbot
        """
        self.verbose = verbose

    def lam_sach_ocr(self, van_ban: str) -> str:
        if not van_ban:
            return ""

        thay_the = {
            'Mó': 'Mỏ', 'chât': 'chất', 'QĐÐ': 'QĐ',
            'đục': 'dục', 'đồi': 'đổi', 'bồ': 'bổ',
            'bồ sung': 'bổ sung', 'sửa đỏi': 'sửa đổi',
            'sửa đồi': 'sửa đổi', 'phề`^1n': 'phần',
            'học phe`^1n': 'học phần',
            'HƯMG': 'HUMG', 'CTCT-SV': 'CTCT-SV',
            'Ô': '0', 'ø': '0', 'Ó': '0', 'ƒ': '2',
            't/hiện': 'thực hiện', 'p/hợp': 'phối hợp',
            'Địa chât': 'Địa chất', 'Mó - Địa chât': 'Mỏ - Địa chất',
        }

        for sai, dung in thay_the.items():
            van_ban = van_ban.replace(sai, dung)

        van_ban = re.sub(r'\s+', ' ', van_ban)
        van_ban = re.sub(
            r'[^\w\s\.,;:!?\-\(\)/đĐáàảãạăắằẳẵặâấầẩẫậéèẻẽẹêếềểễệíìỉĩịóòỏõọôốồổỗộơớờởỡợúùủũụưứừửữựýỳỷỹỵÁÀẢÃẠĂẮẰẲẴẶÂẤẦẨẪẬÉÈẺẼẸÊẾỀỂỄỆÍÌỈĨỊÓÒỎÕỌÔỐỒỔỖỘƠỚỜỞỠỢÚÙỦŨỤƯỨỪỬỮỰÝỲỶỸỴ]',
            ' ',
            van_ban
        )
        return van_ban.strip()

    def read_word(self, path: str):
        docs = []
        try:
            doc = DocxDocument(path)
            name = os.path.basename(path)

            # ===== PARAGRAPH =====
            for idx, p in enumerate(doc.paragraphs):
                text_raw = p.text.strip()
                if not text_raw:
                    continue

                text_clean = self.lam_sach_ocr(text_raw)
                if not text_clean:
                    continue

                docs.append(
                    Document(
                        page_content=f"[Word:{name}][Paragraph {idx + 1}] {text_clean}",
                        metadata={
                            "source": name,
                            "type": "paragraph",
                            "paragraph_index": idx + 1
                        }
                    )
                )

            # ===== TABLE =====
            for t_idx, table in enumerate(doc.tables):
                for r_idx, row in enumerate(table.rows):
                    cells = []
                    for c_idx, cell in enumerate(row.cells):
                        cell_text = self.lam_sach_ocr(cell.text.strip())
                        if cell_text:
                            cells.append(cell_text)

                    if not cells:
                        continue

                    # Row-level document (CHUẨN RAG)
                    docs.append(
                        Document(
                            page_content=(
                                    f"[Table:{t_idx + 1}][Row:{r_idx + 1}][{name}]\n"
                                    + " | ".join(cells)
                            ),
                            metadata={
                                "source": name,
                                "type": "table_row",
                                "table_index": t_idx + 1,
                                "row_index": r_idx + 1,
                                "num_columns": len(cells)
                            }
                        )
                    )

        except Exception as e:
            logger.exception(f"Lỗi đọc file {path}: {e}")

        return docs

    def read_word_folder(self, folder_path: str):
        docs = []

        docx_files = sorted(
            f for f in os.listdir(folder_path)
            if f.lower().endswith(".docx")
        )

        if self.verbose:
            logger.info(
                f"Tìm thấy {len(docx_files)} file .docx trong thư mục {folder_path}"
            )

        for file in docx_files:
            if self.verbose:
                logger.info(f"Đang xử lý: {file}")

            path = os.path.join(folder_path, file)
            docs.extend(self.read_word(path))

        return docs


if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO)
    processor = DocumentProcessor(verbose=True)
    all_docs = processor.read_word_folder(DOCX_DIR)

    print(f"\nTổng số đoạn văn bản sau khi đọc và làm sạch: {len(all_docs)}")
    if all_docs:
        print("Test doc:")
        print(all_docs[0].page_content)
        print("Metadata:", all_docs[0].metadata)
